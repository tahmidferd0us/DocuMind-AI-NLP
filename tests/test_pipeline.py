import os
import sys
import io
from pathlib import Path

# Ensure UTF-8 output encoding on Windows console
if sys.platform.startswith("win"):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# Add documind-nlp to python path
BASE_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE_DIR))

from src.parsers import parse_document, clean_text
from src.summarizer import generate_extractive_summary, generate_abstractive_summary
from src.rag import split_document_pages, VectorStore, RAGEngine
from src.entities import extract_entities, extract_keywords
from src.analytics import compute_document_analytics
from src.exporters import generate_docx_report, generate_pdf_report
from src.evaluation import evaluate_summary

SAMPLE_TEXT = """
Artificial Intelligence in Modern Healthcare
King's Own Institute (KOI) Research Report. Author: Dr. John Smith, Sydney, Australia. Date: March 2026.

Natural Language Processing (NLP) is revolutionising medical documentation and clinical decision support.
By processing electronic health records, clinical notes, and biomedical literature, NLP systems can extract valuable diagnostic insights.
Large Language Models and Retrieval-Augmented Generation (RAG) enable medical practitioners to query vast clinical datasets with grounded citations.

However, challenges remain around model hallucination, patient data privacy, and algorithm transparency.
Automated summarisation models like LexRank and neural transformers help clinicians quickly review patient histories without missing critical contraindications.
Healthcare organizations including the World Health Organization and local Australian health networks are piloting AI tools to streamline administrative workloads.

In conclusion, ethical AI deployment combined with strict grounded verification provides significant benefits for modern healthcare systems.
"""

def test_text_cleaning():
    print("Testing text cleaning...")
    raw = "Hello   world!\u00a0\u200b This is docu-\nment text.\n\n\n\nNew paragraph."
    cleaned = clean_text(raw)
    assert "document" in cleaned
    assert "\u00a0" not in cleaned
    print("[PASS] Text cleaning passed")

def test_document_parsing():
    print("Testing document parser with TXT and synthetic DOCX...")
    # Test TXT
    res_txt = parse_document(SAMPLE_TEXT.encode("utf-8"), "sample.txt")
    assert res_txt["file_type"] == ".txt"
    assert res_txt["total_words"] > 50
    assert len(res_txt["pages"]) == 1
    print("[PASS] TXT parsing passed")

    # Test DOCX creation & parsing
    import docx
    doc = docx.Document()
    doc.add_heading("Medical AI Summary", 0)
    doc.add_paragraph(SAMPLE_TEXT)
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_buf.seek(0)
    
    res_docx = parse_document(docx_buf.getvalue(), "sample.docx")
    assert res_docx["file_type"] == ".docx"
    assert res_docx["total_words"] > 50
    print("[PASS] DOCX parsing passed")

def test_extractive_summarizer():
    print("Testing extractive summarizer (LexRank)...")
    res = generate_extractive_summary(SAMPLE_TEXT, sentence_count=3)
    assert res["sentence_count"] == 3
    assert len(res["sentences"]) == 3
    assert len(res["summary"]) > 20
    print(f"[PASS] Extractive summary passed ({res['sentence_count']} sentences extracted)")

def test_rag_and_vector_store():
    print("Testing RAG chunking, FAISS vector search, and offline retrieval...")
    doc_data = parse_document(SAMPLE_TEXT.encode("utf-8"), "sample.txt")
    chunks = split_document_pages(doc_data["pages"], chunk_size=300, chunk_overlap=50)
    assert len(chunks) >= 2
    
    vs = VectorStore()
    vs.build_index(chunks)
    
    results = vs.search("What are the challenges of AI in healthcare?", top_k=2)
    assert len(results) > 0
    top_chunk, score = results[0]
    assert score > 0.0
    print(f"[PASS] FAISS Vector search passed (top match score: {score:.3f}, snippet: '{top_chunk.content[:60]}...')")

    # Test RAGEngine
    rag = RAGEngine(vs)
    ans = rag.answer_question("Who authored the report and what organization is mentioned?")
    assert len(ans["sources"]) > 0
    print("[PASS] RAG engine query passed")

def test_entities_and_keywords():
    print("Testing NER (spaCy) and Keywords (KeyBERT)...")
    entities = extract_entities(SAMPLE_TEXT)
    assert entities["total_entities_found"] > 0
    
    keywords = extract_keywords(SAMPLE_TEXT, top_n=5)
    assert len(keywords) > 0
    print(f"[PASS] NER and KeyBERT passed (Found {entities['total_entities_found']} entities, {len(keywords)} key phrases)")

def test_analytics():
    print("Testing reading metrics and readability...")
    metrics = compute_document_analytics(SAMPLE_TEXT, "sample.txt")
    assert metrics["word_count"] > 50
    assert metrics["reading_time_min"] > 0
    assert metrics["readability"]["score"] > 0
    print(f"[PASS] Analytics passed (Word count: {metrics['word_count']}, Flesch score: {metrics['readability']['score']})")

def test_exporters():
    print("Testing DOCX and PDF report exporters...")
    doc_info = parse_document(SAMPLE_TEXT.encode("utf-8"), "sample.txt")
    ext_sum = generate_extractive_summary(SAMPLE_TEXT, sentence_count=3)
    analytics = compute_document_analytics(SAMPLE_TEXT, "sample.txt")
    entities = extract_entities(SAMPLE_TEXT)
    keywords = extract_keywords(SAMPLE_TEXT)
    qa_hist = [{"question": "What are challenges?", "answer": "Challenges include hallucination and privacy.", "sources": [{"page": 1}]}]

    docx_bytes = generate_docx_report(
        document_info=doc_info,
        extractive_summary=ext_sum,
        abstractive_summary={"summary": "Sample abstractive text."},
        entities_data=entities,
        keywords_data=keywords,
        analytics_data=analytics,
        qa_history=qa_hist
    )
    assert len(docx_bytes) > 1000

    pdf_bytes = generate_pdf_report(
        document_info=doc_info,
        extractive_summary=ext_sum,
        abstractive_summary={"summary": "Sample abstractive text."},
        entities_data=entities,
        keywords_data=keywords,
        analytics_data=analytics,
        qa_history=qa_hist
    )
    assert len(pdf_bytes) > 1000
    print(f"[PASS] Exporters passed (DOCX: {len(docx_bytes):,} bytes, PDF: {len(pdf_bytes):,} bytes)")

def test_evaluation():
    print("Testing ROUGE and BLEU metric calculation...")
    ref = "Natural Language Processing helps medical document analysis and clinical decisions."
    cand = "NLP assists clinical documentation and decision support in healthcare."
    scores = evaluate_summary(ref, cand)
    assert "rouge1" in scores
    assert "bleu" in scores
    assert scores["rouge1"]["fmeasure"] > 0
    print(f"[PASS] Evaluation passed (ROUGE-1 F1: {scores['rouge1']['fmeasure']}%, BLEU: {scores['bleu']['score']})")

if __name__ == "__main__":
    print("=" * 60)
    print("RUNNING DOCUMIND NLP PIPELINE INTEGRATION TEST SUITE")
    print("=" * 60)
    test_text_cleaning()
    test_document_parsing()
    test_extractive_summarizer()
    test_rag_and_vector_store()
    test_entities_and_keywords()
    test_analytics()
    test_exporters()
    test_evaluation()
    print("=" * 60)
    print("ALL TESTS PASSED SUCCESSFULLY!")
    print("=" * 60)

import io
import re
import unicodedata
from typing import Dict, Any, List, Union
from pathlib import Path

def clean_text(text: str) -> str:
    """Normalises whitespace, fixes encoding artifacts, and removes boilerplate noise."""
    if not text:
        return ""
    
    # Unicode normalization (NFKC decomposes compatibility chars and standardizes)
    text = unicodedata.normalize("NFKC", text)
    
    # Replace non-breaking spaces, zero-width spaces, special whitespace
    text = text.replace("\u00a0", " ").replace("\u200b", "").replace("\ufeff", "")
    
    # Normalize varied bullet points
    text = re.sub(r'[\u2022\u2023\u25E6\u2043\u2219]', '•', text)
    
    # Normalize quotation marks and apostrophes
    text = re.sub(r'[\u2018\u2019\u201A\u201B]', "'", text)
    text = re.sub(r'[\u201C\u201D\u201E\u201F]', '"', text)
    text = re.sub(r'[\u2013\u2014]', '-', text)
    
    # Fix hyphenation at line breaks (e.g., "docu-\nment" -> "document")
    text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
    
    # Normalize multiple newlines (keep max 2 newlines for paragraph breaks)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Normalize consecutive spaces/tabs within lines
    text = re.sub(r'[^\S\n]+', ' ', text)
    
    return text.strip()

def parse_pdf(file_source: Union[bytes, io.BytesIO, str, Path]) -> List[Dict[str, Any]]:
    """Extracts text per page from PDF using pdfplumber with fallback to pypdf."""
    pages = []
    
    # Ensure BytesIO stream
    if isinstance(file_source, bytes):
        stream = io.BytesIO(file_source)
    elif isinstance(file_source, (str, Path)):
        with open(file_source, "rb") as f:
            stream = io.BytesIO(f.read())
    else:
        stream = file_source

    # Try pdfplumber first
    try:
        import pdfplumber
        stream.seek(0)
        with pdfplumber.open(stream) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                extracted = page.extract_text(layout=True) or page.extract_text() or ""
                cleaned = clean_text(extracted)
                if cleaned:
                    pages.append({
                        "page_number": idx,
                        "raw_text": extracted,
                        "cleaned_text": cleaned,
                        "word_count": len(cleaned.split()),
                        "char_count": len(cleaned)
                    })
    except Exception:
        # Fallback to pypdf
        import pypdf
        stream.seek(0)
        reader = pypdf.PdfReader(stream)
        for idx, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text() or ""
            cleaned = clean_text(extracted)
            if cleaned:
                pages.append({
                    "page_number": idx,
                    "raw_text": extracted,
                    "cleaned_text": cleaned,
                    "word_count": len(cleaned.split()),
                    "char_count": len(cleaned)
                })
                
    return pages

def parse_docx(file_source: Union[bytes, io.BytesIO, str, Path]) -> List[Dict[str, Any]]:
    """Extracts paragraphs and tables from a Word DOCX document."""
    import docx
    
    if isinstance(file_source, bytes):
        stream = io.BytesIO(file_source)
    elif isinstance(file_source, (str, Path)):
        with open(file_source, "rb") as f:
            stream = io.BytesIO(f.read())
    else:
        stream = file_source

    stream.seek(0)
    doc = docx.Document(stream)
    
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text.strip())
            
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_raw = "\n\n".join(paragraphs)
    full_cleaned = clean_text(full_raw)
    
    # DOCX does not have fixed static pages; represent as unified document content
    return [{
        "page_number": 1,
        "raw_text": full_raw,
        "cleaned_text": full_cleaned,
        "word_count": len(full_cleaned.split()),
        "char_count": len(full_cleaned)
    }]

def parse_txt(file_source: Union[bytes, io.BytesIO, str, Path]) -> List[Dict[str, Any]]:
    """Extracts plain text with multi-encoding fallback."""
    if isinstance(file_source, (str, Path)):
        with open(file_source, "rb") as f:
            raw_bytes = f.read()
    elif isinstance(file_source, io.BytesIO):
        raw_bytes = file_source.getvalue()
    else:
        raw_bytes = file_source

    # Detect encoding
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252", "iso-8859-1"]
    decoded = None
    for enc in encodings:
        try:
            decoded = raw_bytes.decode(enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
            
    if decoded is None:
        decoded = raw_bytes.decode("utf-8", errors="ignore")

    cleaned = clean_text(decoded)
    return [{
        "page_number": 1,
        "raw_text": decoded,
        "cleaned_text": cleaned,
        "word_count": len(cleaned.split()),
        "char_count": len(cleaned)
    }]

def parse_document(file_source: Union[bytes, io.BytesIO, str, Path], filename: str) -> Dict[str, Any]:
    """
    Unified entry point for document parsing.
    Supports .pdf, .docx, .txt.
    Returns metadata, cleaned full text, and per-page chunks.
    """
    ext = Path(filename).suffix.lower()
    
    if ext == ".pdf":
        pages = parse_pdf(file_source)
    elif ext in [".docx", ".doc"]:
        pages = parse_docx(file_source)
    elif ext in [".txt", ".md", ".log", ".csv"]:
        pages = parse_txt(file_source)
    else:
        raise ValueError(f"Unsupported file format '{ext}'. Supported formats: .pdf, .docx, .txt")

    if not pages:
        pages = [{
            "page_number": 1,
            "raw_text": "",
            "cleaned_text": "",
            "word_count": 0,
            "char_count": 0
        }]

    full_cleaned_text = "\n\n".join(p["cleaned_text"] for p in pages if p["cleaned_text"])
    full_raw_text = "\n\n".join(p["raw_text"] for p in pages if p["raw_text"])
    
    total_words = sum(p["word_count"] for p in pages)
    total_chars = len(full_cleaned_text)

    return {
        "filename": filename,
        "file_type": ext,
        "page_count": len(pages),
        "total_words": total_words,
        "total_characters": total_chars,
        "cleaned_text": full_cleaned_text,
        "raw_text": full_raw_text,
        "pages": pages
    }

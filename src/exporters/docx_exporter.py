import io
from typing import Dict, Any, List, Optional
from datetime import datetime
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets table cell background color."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generate_docx_report(
    document_info: Dict[str, Any],
    extractive_summary: Optional[Dict[str, Any]] = None,
    abstractive_summary: Optional[Dict[str, Any]] = None,
    entities_data: Optional[Dict[str, Any]] = None,
    keywords_data: Optional[List[Dict[str, Any]]] = None,
    analytics_data: Optional[Dict[str, Any]] = None,
    qa_history: Optional[List[Dict[str, Any]]] = None
) -> bytes:
    """Generates a styled, comprehensive DOCX analytical report."""
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_heading("DocuMind AI — Document Intelligence Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(30, 58, 138) # Dark Brand Blue
        run.font.name = "Arial"

    doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')} | Smart NLP Platform (KOI)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("―" * 45).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 1: Document Metadata & Analytics
    doc.add_heading("1. Document Metadata & Reading Analytics", level=1)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    meta_rows = [
        ("File Name", document_info.get("filename", "Unknown")),
        ("Page Count", str(document_info.get("page_count", 1))),
        ("Total Word Count", f"{analytics_data.get('word_count', document_info.get('total_words', 0)):,} words" if analytics_data else "N/A"),
        ("Estimated Reading Time", f"{analytics_data.get('reading_time_min', 0)} min (Speaking: {analytics_data.get('speaking_time_min', 0)} min)" if analytics_data else "N/A"),
        ("Readability Score (Flesch)", f"{analytics_data.get('readability', {}).get('score', 'N/A')} — {analytics_data.get('readability', {}).get('level', '')} ({analytics_data.get('readability', {}).get('grade', '')})" if analytics_data else "N/A")
    ]
    
    for i, (label, val) in enumerate(meta_rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(row.cells[0], "F1F5F9")
        row.cells[1].text = str(val)

    doc.add_paragraph()

    # Section 2: Summaries
    doc.add_heading("2. Dual Document Summaries", level=1)
    
    # Abstractive
    if abstractive_summary and abstractive_summary.get("summary"):
        h2_a = doc.add_heading("A. Abstractive Neural Summary (Synthesis)", level=2)
        h2_a.paragraph_format.space_before = Pt(8)
        p = doc.add_paragraph(abstractive_summary["summary"])
        p.paragraph_format.line_spacing = 1.15
        
    # Extractive
    if extractive_summary and extractive_summary.get("sentences"):
        h2_b = doc.add_heading(f"B. Extractive Summary (LexRank Top {len(extractive_summary['sentences'])} Sentences)", level=2)
        h2_b.paragraph_format.space_before = Pt(8)
        for s in extractive_summary["sentences"]:
            bp = doc.add_paragraph(f"• {s}")
            bp.paragraph_format.left_indent = Inches(0.2)

    doc.add_paragraph()

    # Section 3: Key Topics & Named Entities
    doc.add_heading("3. Key Topics & Named Entities", level=1)
    
    if keywords_data:
        doc.add_heading("Top Key Phrases (KeyBERT)", level=2)
        kw_text = ", ".join([f"{kw['keyword']} ({kw['score']})" for kw in keywords_data[:12]])
        doc.add_paragraph(kw_text)
        
    if entities_data and entities_data.get("entities_by_type"):
        doc.add_heading("Extracted Entities (spaCy NER)", level=2)
        for group, items in entities_data["entities_by_type"].items():
            if items:
                names = ", ".join([f"{it['name']} (x{it['count']})" for it in items[:8]])
                doc.add_paragraph(f"• {group}: {names}")

    doc.add_paragraph()

    # Section 4: Grounded Q&A Interaction History
    if qa_history:
        doc.add_heading("4. Question Answering Interaction History (RAG)", level=1)
        for idx, qa in enumerate(qa_history, start=1):
            q_p = doc.add_paragraph()
            q_run = q_p.add_run(f"Q{idx}: {qa.get('question', '')}")
            q_run.bold = True
            
            a_p = doc.add_paragraph()
            a_p.add_run(f"Answer: {qa.get('answer', '')}")
            a_p.paragraph_format.left_indent = Inches(0.2)
            
            if qa.get("sources"):
                src_list = [f"Page {s.get('page')}" for s in qa.get("sources", [])]
                src_p = doc.add_paragraph()
                src_run = src_p.add_run(f"Citations: {', '.join(set(src_list))}")
                src_run.italic = True
                src_run.font.color.rgb = RGBColor(100, 116, 139)
                src_p.paragraph_format.left_indent = Inches(0.2)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
